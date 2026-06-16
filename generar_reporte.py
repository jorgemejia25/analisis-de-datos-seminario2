"""
Genera reporte_ejecutivo.docx  —  ejecutar después de generar_figuras.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

IMG = "imgs/"
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

def style_run(run, bold=False, size=12):
    run.font.name      = "Arial"
    run.font.size      = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.bold      = bold

def para(text, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=0, space_after=5):
    p   = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    style_run(run, bold=bold, size=size)
    return p

def heading(text, level=1):
    p = para(text, bold=True,
             size=14 if level == 1 else 12,
             space_before=12, space_after=4)
    return p

def caption(text):
    para(text, bold=False, size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

def img(filename, width=Inches(5.5)):
    path = IMG + filename
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        para(f"[Imagen no encontrada: {filename}]")

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            style_run(run, bold=True, size=10)
        hdr[i].paragraphs[0].paragraph_format.space_after = Pt(2)
        if col_widths:
            hdr[i].width = Inches(col_widths[i])
    for ri, row in enumerate(rows):
        cells = t.rows[ri+1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for run in cells[ci].paragraphs[0].runs:
                style_run(run, bold=False, size=10)
            cells[ci].paragraphs[0].paragraph_format.space_after = Pt(2)
            if col_widths:
                cells[ci].width = Inches(col_widths[ci])
    return t

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    style_run(run, size=12)

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
para("")
para("Reporte Ejecutivo", bold=True, size=16,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=36)
para("Análisis de Sensores IoT para Mantenimiento Predictivo",
     bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Industrias Metalmecánicas del Norte",
     size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4)
para("")
para("Estudiante:  Jorge Andrés Mejía",  size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Carnet:       202300376",           size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Curso:        Seminario de Sistemas 2", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Fecha:        16 de junio de 2026", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════════════════
heading("1. Introducción")
para(
    "Industrias Metalmecánicas del Norte opera 12 máquinas críticas (M01–M12) equipadas con "
    "sensores IoT que registran temperatura (°C), presión (PSI), vibración (mm/s) y potencia (kW), "
    "junto con el estado operativo de cada máquina. El objetivo del presente análisis es identificar "
    "patrones en los datos de sensores que permitan anticipar fallos antes de que ocurran, reduciendo "
    "el tiempo de parada no planificado y los costos de mantenimiento correctivo."
)
para(
    "El dataset original contiene 50 registros distribuidos entre las 12 máquinas, recolectados "
    "durante aproximadamente 21 horas. Se identificaron y corrigieron múltiples problemas de calidad "
    "de datos antes de realizar el análisis exploratorio. El dataset limpio final se entrega como "
    "dataset_limpio.csv."
)

heading("Descripción del dataset", level=2)
table(
    ["Campo", "Tipo", "Descripción"],
    [
        ["id_sensor",           "Entero",   "Identificador único del registro"],
        ["fecha_hora",          "Datetime", "Fecha y hora de la lectura del sensor"],
        ["maquina_id",          "Texto",    "Identificador de la máquina (M01–M12)"],
        ["temperatura_c",       "Decimal",  "Temperatura de operación en grados Celsius"],
        ["presion_psi",         "Decimal",  "Presión operativa en PSI"],
        ["vibracion_mm_s",      "Decimal",  "Nivel de vibración en mm/s"],
        ["potencia_kw",         "Decimal",  "Consumo de potencia en kilovatios"],
        ["estado_operativo",    "Texto",    "Estado: OPERATIVO, MANTENIMIENTO o FALLADO"],
        ["tecnico_responsable", "Texto",    "Técnico asignado a la máquina"],
        ["lote_produccion",     "Texto",    "Lote de producción asociado al registro"],
    ],
    col_widths=[1.5, 1.0, 3.5]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. DATA CLEANING
# ══════════════════════════════════════════════════════════════════════════════
heading("2. Problemas encontrados y soluciones aplicadas")
para(
    "Se identificaron nueve categorías de problemas de calidad de datos en el dataset original. "
    "La tabla a continuación documenta cada problema, los registros afectados y la solución aplicada."
)
table(
    ["#", "Columna", "Problema", "Reg.", "Solución aplicada"],
    [
        ["1",  "fecha_hora",          "Timestamps con minutos imposibles (:66:, :76:, :86:, :96:)",  "4",  "Corregidos a 18:46:05 con expresión regular"],
        ["1b", "fecha_hora",          "Años erróneos: 2035 y 2045 en lugar de 2025",                  "2",  "Reemplazados a 2025 con expresión regular"],
        ["2",  "maquina_id",          "Celdas vacías y typo 'M04D' en lugar de 'M04'",               "4",  "Forward-fill + normalización con regex (M\\d{2})"],
        ["3",  "temperatura_c",       "Valor 'null' como texto + 6 outliers extremos (999, -88.5)",  "7",  "Parseado a NaN; imputado con mediana por máquina"],
        ["4",  "presion_psi",         "3 outliers extremos (valor sentinel 999.9)",                   "3",  "Marcados NaN; imputados con mediana por máquina"],
        ["5",  "vibracion_mm_s",      "5 outliers extremos (10 000 000 y -5 555 555)",                "5",  "Marcados NaN; imputados con mediana por máquina"],
        ["6",  "potencia_kw",         "3 outliers extremos (99.9) + 1 valor 0.0001",                  "4",  "Marcados NaN; imputados con mediana por máquina"],
        ["7",  "tecnico_responsable", "Celdas vacías, 'N/A' y valor basura 'zzzzzzzz'",              "4",  "Reemplazados con técnico real de la máquina"],
        ["8",  "lote_produccion",     "Formatos incorrectos: LOTx001 y LOT/001",                      "5",  "Estandarizados a formato LOT-NNN con regex"],
        ["9",  "duplicados",          "Verificación de duplicados semánticos (misma máq + hora)",     "0",  "Sin duplicados encontrados tras ordenamiento"],
    ],
    col_widths=[0.3, 1.3, 2.5, 0.4, 1.9]
)
para("")
img("02_missing.png", width=Inches(5.5))
caption("Figura 1. Registros con errores por columna antes y después de la limpieza.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. DESCRIPCIÓN GENERAL
# ══════════════════════════════════════════════════════════════════════════════
heading("3. Descripción de los Datos")
para(
    "El dataset limpio contiene 50 registros, 12 columnas y cubre 12 máquinas durante "
    "aproximadamente 21 horas de operación. El 72% de los registros corresponde a estado "
    "OPERATIVO, 14% a FALLADO y 14% a MANTENIMIENTO."
)
img("01_descripcion.png", width=Inches(5.5))
caption("Figura 2. Estado operativo (izquierda) y registros por máquina (derecha).")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. VALORES FALTANTES
# ══════════════════════════════════════════════════════════════════════════════
heading("4. Missing Values, Nulos y Vacíos")
para(
    "Tras la limpieza, el dataset conserva un único valor nulo residual en la columna "
    "inspectedby, la cual no es crítica para el análisis. Todas las variables "
    "numéricas y de identificación de máquina están completamente pobladas."
)
table(
    ["Columna", "Problemas antes", "Después", "%"],
    [
        ["temperatura_c",       "7  (texto 'null' + outliers extremos)", "0", "0.0%"],
        ["presion_psi",         "3  (outliers extremos 999.9)",          "0", "0.0%"],
        ["vibracion_mm_s",      "5  (outliers extremos 10M / −5.5M)",    "0", "0.0%"],
        ["potencia_kw",         "4  (outliers extremos 99.9)",           "0", "0.0%"],
        ["maquina_id",          "4  (vacíos + typo M04D)",               "0", "0.0%"],
        ["tecnico_responsable", "4  (vacíos + N/A + basura)",            "0", "0.0%"],
        ["lote_produccion",     "5  (vacíos + formato incorrecto)",      "0", "0.0%"],
        ["inspectedby",         "1  (vacío)",                            "1", "2.0%"],
    ],
    col_widths=[1.7, 2.8, 0.7, 0.6]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. TENDENCIA CENTRAL
# ══════════════════════════════════════════════════════════════════════════════
heading("5. Medidas de Tendencia Central")
table(
    ["Variable", "Media", "Mediana", "Moda", "Desv. Std.", "Mínimo", "Máximo", "Asimetría"],
    [
        ["Temperatura (°C)",  "81.47", "82.30", "75.80", "10.40", "64.50", "96.20", "−0.07"],
        ["Presión (PSI)",     "39.92", "40.35", "38.00",  "5.72", "29.80", "48.10", "−0.12"],
        ["Vibración (mm/s)",   "2.71",  "2.30",  "2.30",  "2.01",  "1.40", "15.80",  "3.93"],
        ["Potencia (kW)",     "14.94", "15.65", "12.30",  "3.22",  "9.50", "19.40", "−0.12"],
    ],
    col_widths=[1.5, 0.7, 0.8, 0.7, 0.8, 0.7, 0.7, 0.8]
)
para("")
img("03_tendencia_central.png", width=Inches(5.5))
caption("Figura 3. Media (rojo), mediana (azul) y moda (verde) por variable. Banda roja = ±1σ.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. HISTOGRAMAS Y DISTRIBUCIÓN
# ══════════════════════════════════════════════════════════════════════════════
heading("6. Histogramas y Distribución de Datos")
para(
    "Temperatura, presión y potencia presentan distribuciones bimodales, reflejando que "
    "las máquinas operan en dos rangos de carga diferenciados. La vibración muestra "
    "fuerte asimetría positiva (asimetría = 3.93): la mayoría de las lecturas son bajas "
    "pero existen picos que se alejan del rango típico."
)
img("04_histogramas.png", width=Inches(5.5))
caption("Figura 4. Histogramas con curva de densidad KDE superpuesta.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. EVALUACIÓN DE OUTLIERS
# ══════════════════════════════════════════════════════════════════════════════
heading("7. Evaluación de Outliers")
para(
    "Se aplicó el método IQR: un valor es outlier si está fuera del intervalo "
    "[Q1 − 1.5·IQR, Q3 + 1.5·IQR]. La vibración conserva un outlier residual de "
    "15.8 mm/s (M03, estado OPERATIVO) que puede representar una anomalía mecánica real."
)
img("05_dist_iqr.png", width=Inches(5.5))
caption("Figura 5. Distribución de densidad con zonas IQR (verde = normal, rojo = outlier).")
doc.add_page_break()
img("06_boxplots.png", width=Inches(5.5))
caption("Figura 6. Caja de bigotes con estadísticas detalladas y outliers identificados.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. MATRIZ Y MAPA DE CORRELACIÓN
# ══════════════════════════════════════════════════════════════════════════════
heading("8. Correlación entre Variables")
para(
    "Temperatura, presión y potencia presentan correlación muy alta entre sí (r > 0.99), "
    "respondiendo al mismo factor de carga. La vibración tiene correlación baja con el "
    "resto (r ≈ 0.30), convirtiéndola en la variable más informativa de forma independiente "
    "para detectar anomalías mecánicas."
)
img("07_matriz_correlacion.png", width=Inches(4.5))
caption("Figura 7. Matriz de correlación de Pearson.")
doc.add_page_break()
img("08_mapa_correlacion.png", width=Inches(5.5))
caption("Figura 8. Mapa de correlación (pairplot) coloreado por estado operativo.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. CONCLUSIONES Y RECOMENDACIONES
# ══════════════════════════════════════════════════════════════════════════════
heading("10. Conclusiones")
for b in [
    "Temperatura, presión y potencia tienen correlación r > 0.99 entre sí: un único indicador de carga puede representar las tres simultáneamente.",
    "La vibración es la variable más independiente y la más útil para detectar anomalías mecánicas incipientes.",
    "Los fallos registrados son abruptos: no están precedidos de incrementos graduales en temperatura ni presión, lo que apunta a fallos eléctricos o mecánicos súbitos.",
    "M03 es la máquina de mayor riesgo: temperatura promedio alta (92°C) y el único outlier residual de vibración (15.8 mm/s).",
    "El dataset limpio quedó con 50 registros, sin valores atípicos extremos en las variables numéricas clave y un único nulo residual en una columna no crítica.",
]:
    bullet(b)

heading("Recomendaciones para mantenimiento predictivo", level=2)
for r in [
    "Implementar una alerta automática cuando vibracion_mm_s supere 5 mm/s como señal temprana de fallo mecánico.",
    "Priorizar revisión preventiva de M03 y M05 dado su nivel de temperatura sostenidamente alto.",
    "Ampliar el dataset a al menos 30 días para construir un modelo de mantenimiento predictivo estadísticamente robusto.",
    "Registrar la causa de cada fallo (eléctrico, mecánico, desgaste) para enriquecer el análisis en ciclos futuros.",
    "Corregir los errores de captura directamente en el firmware de los sensores (valores sentinel 999, timestamps inválidos).",
    "Dado el alto grado de correlación, un modelo con solo temperatura y vibración puede ser suficiente para clasificar el riesgo de fallo.",
]:
    bullet(r)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11. APÉNDICE
# ══════════════════════════════════════════════════════════════════════════════
heading("11. Apéndice — Entregables y código fuente")
table(
    ["Archivo", "Descripción"],
    [
        ["limpieza_datos.ipynb",  "Notebook de Data Cleaning (Componente 1) — 9 problemas documentados"],
        ["eda.ipynb",             "Notebook de EDA (Componente 2) — todas las secciones requeridas"],
        ["generar_reporte.py",    "Script Python que genera este reporte en formato .docx"],
        ["dataset_limpio.csv",    "Dataset limpio listo para modelado predictivo"],
        ["imgs/",                 "Carpeta con 21 figuras individuales de alta resolución (160 dpi)"],
    ],
    col_widths=[2.0, 4.0]
)
para("")
para("[Enlace a repositorio GitHub / Google Colab]", size=12)

doc.save("reporte_ejecutivo.docx")
print("reporte_ejecutivo.docx generado.")
